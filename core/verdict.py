import os
import time
from pathlib import Path
from docx import Document
from tqdm import tqdm
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_pinecone import PineconeVectorStore
from pinecone import Pinecone
from dotenv import load_dotenv

# 환경 변수 로드
load_dotenv()

def get_llm(model="gpt-4o"):
    llm = ChatOpenAI(model=model)  
    return llm

def extract_text_from_docx(docx_path):
    """
    docx 파일에서 텍스트 추출
    """
    try:
        doc = Document(docx_path)
        full_text = []
        
        print(f"문서에서 텍스트 추출 중: {docx_path}")
        
        for paragraph in tqdm(doc.paragraphs, desc="단락 처리"):
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        extracted_text = '\n'.join(full_text)
        print(f"총 {len(extracted_text)} 글자 추출 완료")
        
        return extracted_text
        
    except Exception as e:
        print(f"문서 추출 오류: {e}")
        return ""

def split_text_into_chunks(text):
    """
    텍스트를 청크로 분할
    """
    print("텍스트를 청크로 분할 중...")
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len,
        separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""]
    )
    
    chunks = text_splitter.split_text(text)
    print(f"총 {len(chunks)}개의 청크로 분할 완료")
    
    return chunks

def create_pinecone_index(pc, index_name="criminal-law-db", dimension=1536):
    """
    Pinecone 인덱스 생성 (존재하지 않는 경우)
    """
    try:
        # 기존 인덱스 확인
        existing_indexes = pc.list_indexes()
        index_names = [idx.name for idx in existing_indexes]
        
        if index_name not in index_names:
            print(f"새 Pinecone 인덱스 생성 중: {index_name}")
            pc.create_index(
                name=index_name,
                dimension=dimension,
                metric='cosine',
                spec={
                    'serverless': {
                        'cloud': 'aws',
                        'region': 'us-east-1'
                    }
                }
            )
            
            # 인덱스 생성 완료까지 대기
            print("인덱스 생성 대기 중...")
            time.sleep(10)
            
        else:
            print(f"기존 인덱스 사용: {index_name}")
            
    except Exception as e:
        print(f"인덱스 생성 오류: {e}")

def embed_and_store_chunks(chunks, pinecone_api_key, openai_api_key, index_name="criminal-law-db"):
    """
    텍스트 청크들을 임베딩하여 Pinecone에 저장
    """
    try:
        print("Pinecone 벡터 스토어 초기화 중...")
        
        # OpenAI 임베딩 모델 초기화 (최신 모델 사용)
        embeddings = OpenAIEmbeddings(
            api_key=openai_api_key,
            model="text-embedding-3-small"  # 최신 임베딩 모델
        )
        
        # Pinecone 벡터 스토어 초기화
        vectorstore = PineconeVectorStore(
            index_name=index_name,
            embedding=embeddings,
            pinecone_api_key=pinecone_api_key
        )
        
        print(f"{len(chunks)}개 청크를 임베딩하여 Pinecone에 저장 중...")
        
        # 청크를 문서 객체로 변환
        from langchain_core.documents import Document as LangChainDocument
        documents = []
        for i, chunk in enumerate(chunks):
            documents.append(LangChainDocument(
                page_content=chunk,
                metadata={
                    "source": "criminal-law.docx",
                    "chunk_id": i,
                    "document_type": "criminal_law"
                }
            ))
        
        # 배치 처리로 저장 (Pinecone 제한 고려)
        batch_size = 50  # 배치 크기 줄임
        for i in tqdm(range(0, len(documents), batch_size), desc="벡터 저장"):
            batch = documents[i:i + batch_size]
            
            try:
                vectorstore.add_documents(documents=batch)
                # API 제한 방지를 위한 지연
                time.sleep(2)
            except Exception as batch_error:
                print(f"배치 {i//batch_size + 1} 저장 중 오류: {batch_error}")
                continue
        
        print("모든 청크가 Pinecone에 성공적으로 저장되었습니다!")
        return vectorstore
        
    except Exception as e:
        print(f"임베딩 및 저장 오류: {e}")
        return None

def process_criminal_law_document(docx_path, pinecone_api_key, openai_api_key):
    """
    형법 문서 전체 처리 파이프라인
    """
    print("형법 문서 처리 시작!")
    
    # 1. docx에서 텍스트 추출
    text = extract_text_from_docx(docx_path)
    if not text:
        print("텍스트 추출 실패")
        return None
    
    # 2. 텍스트를 청크로 분할
    chunks = split_text_into_chunks(text)
    if not chunks:
        print("텍스트 분할 실패")
        return None
    
    # 3. Pinecone 인덱스 생성
    pc = Pinecone(api_key=pinecone_api_key)
    create_pinecone_index(pc)
    
    # 4. 임베딩하여 Pinecone에 저장
    vectorstore = embed_and_store_chunks(chunks, pinecone_api_key, openai_api_key)
    
    if vectorstore:
        print("형법 문서 처리 완료!")
    else:
        print("형법 문서 처리 실패")
        
    return vectorstore

def auto_setup_criminal_law_db():
    """
    형법 데이터베이스 자동 설정 - 필요시에만 실행
    """
    try:
        # API 키 확인
        pinecone_api_key = os.getenv("PINECONE_API_KEY")
        openai_api_key = os.getenv("OPENAI_API_KEY")
        
        if not pinecone_api_key or not openai_api_key:
            print("API 키가 설정되지 않았습니다. 형법 참조 없이 판결합니다.")
            return False
        
        # Pinecone 인덱스 존재 확인
        try:
            pc = Pinecone(api_key=pinecone_api_key)
            existing_indexes = pc.list_indexes()
            index_names = [idx.name for idx in existing_indexes]
            
            if "criminal-law-db" in index_names:
                print("형법 데이터베이스가 이미 존재합니다.")
                return True
        except Exception as e:
            print(f"Pinecone 연결 확인 중 오류: {e}")
        
        # 형법 문서 파일 확인
        docx_path = Path(__file__).parent / "assets" / "criminal-law.docx"
        if not docx_path.exists():
            print("형법 문서 파일을 찾을 수 없습니다. 형법 참조 없이 판결합니다.")
            return False
        
        print("형법 데이터베이스를 자동 설정합니다...")
        
        # 형법 문서 처리 및 저장
        vectorstore = process_criminal_law_document(docx_path, pinecone_api_key, openai_api_key)
        
        if vectorstore:
            print("형법 데이터베이스 자동 설정 완료!")
            return True
        else:
            print("형법 데이터베이스 설정 실패. 형법 참조 없이 판결합니다.")
            return False
            
    except Exception as e:
        print(f"자동 설정 중 오류 발생: {e}")
        print("형법 참조 없이 판결합니다.")
        return False

def get_criminal_law_context(query, top_k=3):
    """
    형법 문서에서 관련 내용을 검색하여 컨텍스트 제공
    """
    try:
        # API 키 확인
        pinecone_api_key = os.getenv("PINECONE_API_KEY")
        openai_api_key = os.getenv("OPENAI_API_KEY")
        
        if not pinecone_api_key or not openai_api_key:
            return ""
        
        # 임베딩 모델 초기화 (최신 모델 사용)
        embeddings = OpenAIEmbeddings(
            api_key=openai_api_key,
            model="text-embedding-3-small"
        )
        
        # Pinecone 벡터 스토어 연결
        vectorstore = PineconeVectorStore(
            index_name="criminal-law-db",
            embedding=embeddings,
            pinecone_api_key=pinecone_api_key
        )
        
        # 유사도 검색 수행
        docs = vectorstore.similarity_search(query, k=top_k)
        
        if not docs:
            return ""
        
        # 검색된 문서들을 텍스트로 결합
        context_parts = []
        for i, doc in enumerate(docs, 1):
            context_parts.append(f"[형법 조항 {i}]\n{doc.page_content}\n")
        
        context = "\n".join(context_parts)
        print(f"{len(docs)}개의 관련 형법 조항을 찾았습니다.")
        
        return context
        
    except Exception as e:
        print(f"형법 검색 중 오류: {e}")
        return ""

def create_search_query_from_messages(message_list):
    """
    대화 내용에서 형법 검색을 위한 쿼리 생성
    """
    # 검사와 변호사의 주장에서 핵심 키워드 추출
    key_terms = []
    
    for message in message_list:
        if message['role'] in ['검사', '변호사']:
            content = message['content']
            # 법적 용어나 범죄 관련 키워드 우선 추출
            if any(term in content for term in ['살인', '절도', '사기', '폭행', '강도', '횡령', '배임', '방화', '강간', '유괴']):
                key_terms.append(content)
    
    # 검색 쿼리 생성
    if key_terms:
        search_query = " ".join(key_terms[:2])  # 처음 2개 메시지만 사용
    else:
        # 모든 메시지 결합
        all_content = " ".join([msg['content'] for msg in message_list if msg['role'] in ['검사', '변호사']])
        search_query = all_content[:500]  # 길이 제한
    
    return search_query

def get_judge_result(message_list):
    """
    형법을 참고하여 판결 결과 생성
    """
    # 형법 데이터베이스 자동 설정 (필요시에만)
    db_ready = auto_setup_criminal_law_db()
    
    llm = get_llm()
    
    # 형법 검색 및 컨텍스트 생성
    criminal_law_context = ""
    if db_ready:
        search_query = create_search_query_from_messages(message_list)
        print(f"형법 검색 쿼리: {search_query[:100]}...")
        criminal_law_context = get_criminal_law_context(search_query)
    
    # 프롬프트 템플릿 (형법 참조 포함 또는 기본)
    if criminal_law_context:
        prompt = ChatPromptTemplate.from_template("""
            당신은 공정한 AI 판사입니다. 아래는 검사와 변호사의 주장 및 증거입니다. 
            제공된 형법 조항을 참고하여 누가 더 논리적이고 설득력 있는 주장을 했는지 판단해주세요.
            
            [관련 형법 조항]:
            {criminal_law_context}
            
            [전체 대화 내용]:
            {messages}
            
            위의 형법 조항들을 참고하여 법적 근거를 바탕으로 판단하고, 
            어떤 형법 조항이 이 사건에 적용되는지 명시해주세요.

            판결 형식:
            [승자]: (검사 또는 변호사)
            [적용 법조항]: (관련 형법 조항 번호 및 내용 요약)
            [판결 이유]: (형법을 근거로 한 상세하고 논리적인 이유)
        """)
    else:
        print("형법 조항을 참조할 수 없어 일반적인 판결을 진행합니다.")
        prompt = ChatPromptTemplate.from_template("""
            당신은 공정한 AI 판사입니다. 아래는 검사와 변호사의 주장 및 증거입니다. 
            이 정보를 바탕으로 누가 더 논리적이고 설득력 있는 주장을 했는지 판단해주세요.
            그 이유도 간단히 설명해주세요.

            [전체 대화 내용]:
            {messages}

            판결 형식:
            [승자]: (검사 또는 변호사)
            [판결 이유]: (간단하고 논리적인 이유)
        """)

    # 메시지 포맷팅
    messages_joined = "\n".join([f"[{m['role']}]: {m['content']}" for m in message_list if m['role'] in ['검사', '변호사']])
    
    # 체인 생성 및 실행
    chain = prompt | llm | StrOutputParser()
    
    if criminal_law_context:
        result = chain.invoke({
            "messages": messages_joined,
            "criminal_law_context": criminal_law_context
        })
    else:
        result = chain.invoke({"messages": messages_joined})
    
    return result

# 테스트 사례 추가
def get_test_case_messages():
    """
    테스트용 검사/변호사 대화 내용 반환
    """
    return [
        {
            "role": "검사",
            "content": "피고인 김소현은 2024년 11월 15일 오후 11시경, 자신이 운영하는 바에서 피해자 최지훈에게 독성 물질이 포함된 술을 제공하여 의도적으로 살해한 혐의를 받고 있습니다. 목격자 남기효의 증언에 따르면, 피고인이 피해자에게만 특별히 제조한 칵테일을 제공했으며, 피해자는 이를 마신 후 즉시 중독 증상을 보이며 사망했습니다. 이는 명백한 고의적 살인행위로 형법 제250조 제1항에 해당됩니다."
        },
        {
            "role": "변호사",
            "content": "검사 측 주장은 추측에 불과합니다. 우리 의뢰인 김소현은 단순히 평소처럼 술을 제조하여 제공했을 뿐이며, 피해자의 사망은 예측할 수 없었던 알레르기 반응이나 기존 질병에 의한 것일 가능성이 높습니다. 참고인 우민영의 증언에 따르면, 피고인은 평소 손님들에게 친절하게 서비스를 제공했으며, 피해자와 개인적인 원한 관계도 없었습니다. 고의성이 입증되지 않은 상황에서 살인죄를 적용하는 것은 부당합니다."
        },
        {
            "role": "검사",
            "content": "변호사 측의 주장은 사실과 다릅니다. 부검 결과 피해자의 혈액에서 일반적으로 술에 포함되지 않는 독성 화학물질이 검출되었으며, 이는 의도적으로 투입된 것으로 보입니다. 또한 CCTV 분석 결과, 피고인이 피해자에게만 별도의 재료를 추가하는 모습이 명확히 포착되었습니다. 이러한 객관적 증거들은 피고인의 고의적 살인 의도를 분명히 보여줍니다."
        },
        {
            "role": "변호사",
            "content": "검사 측이 제시한 부검 결과와 CCTV 영상에 대해 이의를 제기합니다. 첫째, 해당 화학물질은 바 청소용 세제에서 유래된 것으로 보이며, 우발적 오염 가능성을 배제할 수 없습니다. 둘째, CCTV 영상은 화질이 불분명하여 정확히 무엇을 첨가했는지 식별이 불가능합니다. 피고인은 평소 칵테일에 특별한 가니시나 시럽을 추가하는 습관이 있었으며, 이는 단순한 서비스 차원의 행동이었습니다. 합리적 의심을 넘어서는 입증이 부족한 상황입니다."
        }
    ]

# 테스트 실행 함수
def run_test_case():
    """
    테스트 사례로 RAG 시스템 검증
    """
    print("=== RAG 시스템 테스트 시작 ===")
    test_messages = get_test_case_messages()
    
    print("\n[테스트 사건 배경]")
    print("- 피고인: 김소현 (바 운영자)")
    print("- 피해자: 최지훈 (고객)")
    print("- 목격자: 남기효")
    print("- 참고인: 우민영")
    print("- 사건: 바에서 독성 물질 투입으로 인한 사망 사건")
    
    print("\n[검사-변호사 주장 내용]")
    for i, msg in enumerate(test_messages, 1):
        print(f"{i}. [{msg['role']}]: {msg['content'][:100]}...")
    
    print("\n[AI 판사 판결 진행중...]")
    result = get_judge_result(test_messages)
    
    print("\n=== AI 판사 최종 판결 ===")
    print(result)
    print("\n=== 테스트 완료 ===")

if __name__ == "__main__":
    # 테스트 실행
    run_test_case()